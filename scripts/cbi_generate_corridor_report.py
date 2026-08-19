"""
Generates one Word document per corridor summarizing recurring bottlenecks
and their characterization — the automated, per-corridor equivalent of
Part II of the CBI technical report.

Requires: pip install python-docx

Reads from the same views/tables the technical report documents:
  - "Year_2025".vw_bottleneck_dashboard_ranked  (severity ranking, Section 11)
  - "Year_2025".segment_recurring_bottlenecks   (bottleneck locations, 9.1)
  - "Year_2025".bottleneck_daily_metrics        (characterization, Section 10)
"""

from __future__ import annotations

import tempfile
from datetime import date
from pathlib import Path

import polars as pl
import psycopg
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import cbi_database
import cbi_map_geometry
from cbi_corridor_registry import CorridorContext

NAVY = RGBColor(0x1F, 0x38, 0x64)


def _load_ranked_bottlenecks(corridor: CorridorContext) -> pl.DataFrame:
    query = f"""
        SELECT
            bottleneck_id, corridor_severity_rank AS severity_rank,
            representative_intersection,
            occurrence_pct, analyzed_days, occurrence_days,
            avg_active_minutes, median_active_minutes, p95_active_minutes,
            annual_queue_mile_hours, avg_max_queue_miles, p95_max_queue_miles,
            annual_max_queue_miles,
            avg_congested_speed_ratio, avg_queue_growth_rate_mph,
            avg_queue_dissipation_rate_mph, severity_index
        FROM "Year_2025".vw_bottleneck_dashboard_ranked
        WHERE corridor = '{corridor.road}'
          AND direction = '{corridor.direction}'
        ORDER BY corridor_severity_rank
    """
    return pl.read_database_uri(
        query=query, uri=cbi_database.database_uri(), engine="connectorx"
    )


def _load_bottleneck_reference_speeds(corridor: CorridorContext) -> pl.DataFrame:
    """Average reference (free-flow) speed observed across each bottleneck's
    segment range, used only to project the stored congested-speed RATIO
    into an approximate mph figure — reference speed varies by time of day
    in the source data, so this is an approximation, not a posted speed
    limit or a single fixed value."""
    query = f"""
        WITH bottleneck_segments AS (
            SELECT b.bottleneck_id, cs.tmc
            FROM "Year_2025".segment_recurring_bottlenecks AS b
            JOIN "Year_2025".corridor_segments AS cs
              ON cs.corridor_id = {corridor.corridor_id}
             AND cs.segment_order BETWEEN b.start_segment_order AND b.end_segment_order
            WHERE b.corridor = '{corridor.road}' AND b.direction = '{corridor.direction}'
        )
        SELECT bs.bottleneck_id, ROUND(AVG(r.reference_speed)::numeric, 1) AS avg_reference_speed_mph
        FROM bottleneck_segments AS bs
        JOIN "Year_2025".probe_readings AS r ON r.tmc_code = bs.tmc
        WHERE r.reference_speed > 0
        GROUP BY bs.bottleneck_id
    """
    return pl.read_database_uri(
        query=query, uri=cbi_database.database_uri(), engine="connectorx"
    )


def _load_bottleneck_daily_detail(corridor: CorridorContext) -> pl.DataFrame:
    """Day-by-day characterization for every occurrence day of every
    bottleneck on this corridor — the granular record behind the summary
    statistics, for inspecting how the detection methodology characterizes
    each individual occurrence (onset/peak/clearance time, queue extent,
    speed). Limited to days the bottleneck actually occurred; non-occurrence
    days carry no onset/peak/queue data and aren't informative for pattern
    inspection."""
    query = f"""
        SELECT
            dm.bottleneck_id, dm.analysis_date, dm.onset_time, dm.peak_time,
            dm.clearance_time, dm.active_congestion_minutes,
            dm.maximum_queue_miles, dm.average_congested_speed_ratio
        FROM "Year_2025".bottleneck_daily_metrics AS dm
        JOIN "Year_2025".segment_recurring_bottlenecks AS b
          ON b.bottleneck_id = dm.bottleneck_id
        WHERE b.corridor = '{corridor.road}' AND b.direction = '{corridor.direction}'
          AND dm.occurrence
        ORDER BY dm.bottleneck_id, dm.analysis_date
    """
    return pl.read_database_uri(
        query=query, uri=cbi_database.database_uri(), engine="connectorx"
    )


def _load_corridor_summary(corridor: CorridorContext) -> dict:
    query = f"""
        SELECT
            COUNT(DISTINCT c.tmc) AS segments,
            ROUND(SUM(c.miles)::numeric, 2) AS corridor_miles,
            (
                SELECT COUNT(DISTINCT analysis_date)
                FROM "Year_2025".congestion_events
                WHERE corridor = '{corridor.road}'
                  AND direction = '{corridor.direction}'
            ) AS analyzed_days,
            (
                SELECT COUNT(*)
                FROM "Year_2025".congestion_events
                WHERE corridor = '{corridor.road}'
                  AND direction = '{corridor.direction}'
            ) AS total_events
        FROM "Year_2025".corridor_segments AS c
        WHERE c.corridor_id = {corridor.corridor_id}
    """
    result = pl.read_database_uri(
        query=query, uri=cbi_database.database_uri(), engine="connectorx"
    )
    return result.row(0, named=True)


def _build_corridor_static_map(corridor: CorridorContext, tmp_dir: Path) -> Path | None:
    """Static PNG of this corridor's bottleneck segments — same real-road
    geometry, directional offset, and Google Maps-style severity color ramp
    as the master HTML report's Region Map tab (see cbi_map_geometry.py).
    Colored against the SAME region-wide K-means fit used everywhere else
    in the report, not refit on just this corridor's handful of
    bottlenecks (which would rescale the color ramp locally and break "one
    color means the same severity everywhere"). Returns None if this
    corridor has no ranking-eligible bottleneck segments to show."""
    all_segments = cbi_map_geometry.load_bottleneck_map_segments()
    all_records, _ = cbi_map_geometry.bottleneck_records_and_legend(all_segments)
    corridor_records = [
        record[:3] for record in all_records if record[3] == corridor.corridor_id
    ]
    if not corridor_records:
        return None

    png_path = tmp_dir / f"{corridor.slug}_map.png"
    cbi_map_geometry.render_static_map_png(corridor_records, png_path)
    return png_path


def _add_heading(doc: Document, text: str, size: int = 16) -> None:
    heading = doc.add_heading(level=1)
    run = heading.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = NAVY


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)


def _fmt(value, decimals: int = 1) -> str:
    if value is None:
        return "\u2013"
    try:
        return f"{float(value):,.{decimals}f}"
    except (TypeError, ValueError):
        return str(value)


def _fmt_time(value) -> str:
    if value is None:
        return "\u2013"
    try:
        return value.strftime("%H:%M")
    except AttributeError:
        return str(value)


def _set_landscape(doc: Document) -> None:
    """This report is dense with methodology-inspection tables (up to ~16
    columns in places) \u2014 landscape gives them room to be read without
    wrapping, rather than confining the whole document to portrait for the
    sake of the shorter summary tables."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def _add_table(doc: Document, header_labels: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header_labels))
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    for i, label in enumerate(header_labels):
        header[i].text = label
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)


def generate_report(
    connection: psycopg.Connection,
    corridor: CorridorContext,
    output_root: Path,
) -> Path:
    ranked = _load_ranked_bottlenecks(corridor)
    summary = _load_corridor_summary(corridor)

    if not ranked.is_empty():
        ref_speeds = _load_bottleneck_reference_speeds(corridor)
        ranked = ranked.join(ref_speeds, on="bottleneck_id", how="left").with_columns(
            (pl.col("avg_congested_speed_ratio") * pl.col("avg_reference_speed_mph"))
            .round(1)
            .alias("avg_speed_mph")
        )
        daily_detail = _load_bottleneck_daily_detail(corridor)
        ref_speed_by_bottleneck = {
            bid: (float(speed) if speed is not None else None)
            for bid, speed in zip(
                ranked["bottleneck_id"].to_list(), ranked["avg_reference_speed_mph"].to_list()
            )
        }
    else:
        daily_detail = pl.DataFrame()
        ref_speed_by_bottleneck = {}

    doc = Document()
    _set_landscape(doc)

    title = doc.add_heading(level=0)
    title_run = title.add_run(f"{corridor.corridor_name}")
    title_run.font.color.rgb = NAVY
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle.add_run(
        "Congestion and Bottleneck Identification \u2014 Automated Corridor Report"
    )
    subtitle_run.italic = True
    doc.add_paragraph(f"Generated: {date.today().isoformat()}")

    _add_heading(doc, "Corridor Summary")
    _add_kv_table(
        doc,
        [
            ("Corridor", corridor.corridor_name),
            ("Segments", summary["segments"]),
            ("Corridor Length (miles)", summary["corridor_miles"]),
            ("Days Analyzed", summary["analyzed_days"]),
            ("Total Congestion Events", summary["total_events"]),
            ("Recurring Bottlenecks Identified", ranked.height),
        ],
    )

    with tempfile.TemporaryDirectory() as tmp_dir_str:
        map_png = _build_corridor_static_map(corridor, Path(tmp_dir_str))
        if map_png:
            _add_heading(doc, "Corridor Map")
            doc.add_paragraph(
                "Colored sections are the exact extent of a detected recurring "
                "bottleneck (not the whole corridor), colored by AADT-weighted "
                "severity on the same Google Maps-style traffic scale (light "
                "orange = least severe of these confirmed bottlenecks, dark "
                "maroon = most severe) used throughout this report — see the "
                "master HTML report's Region Map tab for the interactive version."
            )
            doc.add_picture(str(map_png), width=Inches(9))

    if ranked.is_empty():
        doc.add_paragraph(
            "No recurring bottlenecks met the detection thresholds for "
            "this corridor. This may mean the corridor is genuinely "
            "uncongested, or that the pipeline has not yet completed all "
            "stages for it \u2014 check corridor_pipeline_runs."
        )
    else:
        _add_heading(doc, "Recurring Bottlenecks \u2014 Ranked by Severity")
        doc.add_paragraph(
            "Severity combines annual queue mile-hours, occurrence frequency, "
            "and average congested speed ratio (see the CBI technical report, "
            "Section 11, for the exact formula). Avg Speed (mph) is an "
            "approximation: the stored congested-speed ratio times this "
            "bottleneck's average observed reference (free-flow) speed, which "
            "varies by time of day in the source data \u2014 not a posted speed limit."
        )

        _add_table(
            doc,
            [
                "Rank", "Location", "Occurrence %", "Avg Active Min",
                "Avg Max Queue (mi)", "P95 Max Queue (mi)", "Avg Speed (mph)",
                "Avg Speed Ratio", "Annual Mile-Hrs", "Severity",
            ],
            [
                [
                    row["severity_rank"],
                    row["representative_intersection"] or "",
                    f"{_fmt(row['occurrence_pct'])}%",
                    _fmt(row["avg_active_minutes"]),
                    _fmt(row["avg_max_queue_miles"], 2),
                    _fmt(row["p95_max_queue_miles"], 2),
                    _fmt(row["avg_speed_mph"]),
                    _fmt(row["avg_congested_speed_ratio"], 3),
                    _fmt(row["annual_queue_mile_hours"]),
                    _fmt(row["severity_index"]),
                ]
                for row in ranked.iter_rows(named=True)
            ],
        )

        _add_heading(doc, "Bottleneck Detail")
        for row in ranked.iter_rows(named=True):
            doc.add_heading(
                f"#{row['severity_rank']} \u2014 {row['representative_intersection']}",
                level=2,
            )
            _add_kv_table(
                doc,
                [
                    ("Occurrence", f"{_fmt(row['occurrence_pct'])}% of analyzed days"),
                    ("Occurrence Days", f"{row['occurrence_days']} of {row['analyzed_days']}"),
                    ("Avg / Median / P95 Active Minutes",
                     f"{_fmt(row['avg_active_minutes'])} / "
                     f"{_fmt(row['median_active_minutes'])} / {_fmt(row['p95_active_minutes'])}"),
                    ("Average Max Queue (mi)", _fmt(row["avg_max_queue_miles"], 2)),
                    ("95th Percentile Max Queue (mi)", _fmt(row["p95_max_queue_miles"], 2)),
                    ("Annual Max Queue (mi)", _fmt(row["annual_max_queue_miles"], 2)),
                    ("Average Speed (mph, approx.)", _fmt(row["avg_speed_mph"])),
                    ("Average Congested Speed Ratio", _fmt(row["avg_congested_speed_ratio"], 3)),
                    ("Queue Growth Rate (mph)", _fmt(row["avg_queue_growth_rate_mph"], 2)),
                    ("Queue Dissipation Rate (mph)", _fmt(row["avg_queue_dissipation_rate_mph"], 2)),
                    ("Annual Queue Mile-Hours", _fmt(row["annual_queue_mile_hours"])),
                    ("Severity Index", _fmt(row["severity_index"])),
                ],
            )

            bottleneck_days = daily_detail.filter(pl.col("bottleneck_id") == row["bottleneck_id"])
            if not bottleneck_days.is_empty():
                ref_speed = ref_speed_by_bottleneck.get(row["bottleneck_id"])
                doc.add_heading("Daily Occurrence Detail", level=3)
                doc.add_paragraph(
                    f"Every day this bottleneck occurred ({bottleneck_days.height} of "
                    f"{row['analyzed_days']} analyzed days) \u2014 the underlying daily "
                    "record the summary statistics above are computed from."
                )
                _add_table(
                    doc,
                    [
                        "Date", "Onset", "Peak", "Clearance", "Active Min",
                        "Max Queue (mi)", "Avg Speed (mph)", "Avg Speed Ratio",
                    ],
                    [
                        [
                            d["analysis_date"].isoformat(),
                            _fmt_time(d["onset_time"]),
                            _fmt_time(d["peak_time"]),
                            _fmt_time(d["clearance_time"]),
                            _fmt(d["active_congestion_minutes"], 0),
                            _fmt(d["maximum_queue_miles"], 2),
                            _fmt(
                                d["average_congested_speed_ratio"] * ref_speed
                                if ref_speed is not None and d["average_congested_speed_ratio"] is not None
                                else None
                            ),
                            _fmt(d["average_congested_speed_ratio"], 3),
                        ]
                        for d in bottleneck_days.iter_rows(named=True)
                    ],
                )

    output_dir = output_root / corridor.slug
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{corridor.slug}_congestion_report_{date.today().isoformat()}.docx"
    doc.save(output_path)

    return output_path
